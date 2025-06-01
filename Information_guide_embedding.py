
import random
import string
import os
import fitz  # PyMuPDF
from typing import List, Dict
import openai
import pytesseract
from PIL import Image
import io
from dotenv import load_dotenv
from pinecone import Pinecone, ServerlessSpec
from docx import Document
import re

load_dotenv()

class WordVectorizerOpenAIPinecone:
    def __init__(
        self,
        folder_path: str,
        pinecone_index_name: str,
    ):
        """
        Initialize the PDF vectorizer with OpenAI embeddings and Pinecone storage.
        Processes PDF files from the specified folder path, using the PDF name as part of the vector ID.
        """
        if not os.path.exists(folder_path):
            raise FileNotFoundError(f"Folder not found at: {folder_path}")
        self.folder_path = folder_path
        self.index_name = pinecone_index_name

        # OpenAI setup
        self.embedding_model = "text-embedding-ada-002"
        openai.api_key = os.getenv("OPENAI_API_KEY")

        # Pinecone setup
        self.pc = Pinecone(os.environ.get('PINECONE_API_KEY'))

        if pinecone_index_name not in [i.name for i in self.pc.list_indexes()]:
            print(f"Creating new index: {pinecone_index_name}")
            self.pc.create_index(
                name=pinecone_index_name,
                dimension=1536,  # Dimension for text-embedding-3-small
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1")
            )

        self.index = self.pc.Index(pinecone_index_name)

    def generate_unique_id(self) -> str:
        """Generate a 5-character unique alphanumeric ID."""
        return ''.join(random.choices(string.ascii_lowercase + string.digits, k=5))

    def sanitize_vector_id(self, vector_id: str) -> str:
        """Sanitize the vector ID to ensure it contains only ASCII characters and is valid for Pinecone.
        Replaces spaces, parentheses, and other non-alphanumeric characters with underscores."""
        # Replace spaces, parentheses, and other special characters with underscores
        sanitized = re.sub(r'[^\w\d]', '_', vector_id)
        # Ensure it's ASCII only by replacing any remaining non-ASCII characters
        sanitized = re.sub(r'[^\x00-\x7F]+', '_', sanitized)
        # Truncate if too long (Pinecone has limits on ID length)
        if len(sanitized) > 64:
            # Keep a prefix and add a hash-like suffix for uniqueness
            prefix = sanitized[:54]  # Leave room for unique suffix
            suffix = ''.join(random.choices(string.ascii_lowercase + string.digits, k=10))
            sanitized = f"{prefix}_{suffix}"
        return sanitized
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file."""
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n\n".join(text_content)
        
        except Exception as e:
            print(f"Error extracting text from DOCX {file_path}: {e}")
            return ""

    def extract_text_from_pdf_page(self, doc: fitz.Document, page_num: int) -> str:
        """Extract text from a single PDF page, including OCR from images."""
        try:
            page = doc.load_page(page_num)
            text_content = page.get_text().strip()

            # Perform OCR on the page as an image
            ocr_texts = []
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            try:
                ocr_text = pytesseract.image_to_string(img).strip()
                if ocr_text:
                    ocr_texts.append(ocr_text)
            except Exception as e:
                print(f"OCR error on page {page_num + 1}: {e}")

            # Process embedded images
            image_list = page.get_images(full=True)
            for img_idx, img_info in enumerate(image_list):
                xref = img_info[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image = Image.open(io.BytesIO(image_bytes))
                try:
                    ocr_text = pytesseract.image_to_string(image).strip()
                    if ocr_text:
                        ocr_texts.append(ocr_text)
                except Exception as e:
                    print(f"OCR error on image {img_idx} of page {page_num + 1}: {e}")

            # Combine page text and OCR results
            full_text = text_content + "\n\n" + "\n\n".join(ocr_texts)
            return full_text.strip()

        except Exception as e:
            print(f"Error extracting text from page {page_num + 1}: {e}")
            return ""

    def create_embedding(self, text: str) -> List[float]:
        """Create an embedding for the given text using OpenAI's embedding model."""
        try:
            # Truncate text to avoid exceeding token limits (approx. 8192 tokens for text-embedding-3-small)
            truncated = text[:8000]
            response = openai.Embedding.create(
                model=self.embedding_model,
                input=truncated
            )
            return response['data'][0]['embedding']
        except Exception as e:
            print(f"Error creating embedding: {e}")
            return []

    def delete_vectors_by_document_name(self, document_name: str):
        """Delete all existing vectors in the index related to a specific document name."""
        try:
            results = self.index.query(
                vector=[0.0] * 1536,
                top_k=10000,
                include_metadata=True,
                filter={"document_name": {"$eq": document_name}}
            )
            
            if results and "matches" in results:
                vector_ids = [match["id"] for match in results["matches"] if "id" in match]
                
                if vector_ids:
                    batch_size = 100
                    for i in range(0, len(vector_ids), batch_size):
                        batch = vector_ids[i:i+batch_size]
                        self.index.delete(ids=batch)
                    
                    print(f"Deleted {len(vector_ids)} vectors for document: {document_name}")
                else:
                    print(f"No vectors found for document: {document_name}")
            else:
                print(f"No matches found for document: {document_name}")
                
        except Exception as e:
            print(f"Error deleting vectors for document {document_name}: {e}")



    def embed_and_store_document(self, file_path: str, title: str, replace_existing: bool = True):
        """
        Process a single PDF or DOCX document and create embeddings.
        Store embeddings in Pinecone with a unique ID.
        If replace_existing is True, delete existing vectors for the same document before adding new ones.
        """
        file_name = os.path.splitext(os.path.basename(file_path))[0]  # Extract file name without extension
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            # If replacing existing vectors, delete them first
            if replace_existing:
                self.delete_vectors_by_document_name(file_name)  # This method works for any document type
            
            # Extract text based on file type
            if file_extension == '.pdf':
                doc = fitz.open(file_path)
                print(f"Processing PDF {file_path} with {len(doc)} pages")
                
                # List to collect vectors for batch upsert
                vectors_batch = []
                batch_size = 50
                
                for page_num in range(len(doc)):
                    # Extract text from the page
                    page_text = self.extract_text_from_pdf_page(doc, page_num)
                    
                    if not page_text:
                        continue

                    # Create embedding for the page
                    embedding = self.create_embedding(page_text)
                    if not embedding:
                        continue

                    # Generate vector ID
                    sanitized_file_name = self.sanitize_vector_id(file_name)
                    unique_id = self.generate_unique_id()
                    vector_id = f"{sanitized_file_name}_{page_num}_{unique_id}"
                    vector_id = self.sanitize_vector_id(vector_id)

                    # Create metadata
                    metadata = {
                        "file_name": os.path.basename(file_path),
                        "document_name": file_name,
                        "title": title,
                        "page_number": page_num + 1,
                        "text": page_text[:8000],
                        "char_count": len(page_text),
                        "file_type": "pdf"
                    }

                    vectors_batch.append({
                        "id": vector_id,
                        "values": embedding,
                        "metadata": metadata
                    })
                    
                    if len(vectors_batch) >= batch_size or page_num == len(doc) - 1:
                        if vectors_batch:
                            self.index.upsert(vectors=vectors_batch)
                            print(f"Uploaded batch of {len(vectors_batch)} vectors")
                            vectors_batch = []

                doc.close()
                
            elif file_extension == '.docx':
                print(f"Processing DOCX {file_path}")
                
                # Extract all text from DOCX
                document_text = self.extract_text_from_docx(file_path)
                
                if not document_text:
                    print(f"Warning: No text extracted from {file_path}")
                    return

                # For DOCX, we'll chunk the text into smaller pieces since there are no natural pages
                # Split by double line breaks or every 2000 characters, whichever comes first
                text_chunks = []
                paragraphs = document_text.split('\n\n')
                
                current_chunk = ""
                for paragraph in paragraphs:
                    if len(current_chunk) + len(paragraph) > 2000 and current_chunk:
                        text_chunks.append(current_chunk.strip())
                        current_chunk = paragraph
                    else:
                        current_chunk += "\n\n" + paragraph if current_chunk else paragraph
                
                if current_chunk:
                    text_chunks.append(current_chunk.strip())
                
                # If no good chunks were created, just use the whole text
                if not text_chunks:
                    text_chunks = [document_text]
                
                vectors_batch = []
                batch_size = 50
                
                for chunk_num, chunk_text in enumerate(text_chunks):
                    # Create embedding for the chunk
                    embedding = self.create_embedding(chunk_text)
                    if not embedding:
                        continue

                    # Generate vector ID
                    sanitized_file_name = self.sanitize_vector_id(file_name)
                    unique_id = self.generate_unique_id()
                    vector_id = f"{sanitized_file_name}_{chunk_num}_{unique_id}"
                    vector_id = self.sanitize_vector_id(vector_id)

                    # Create metadata
                    metadata = {
                        "file_name": os.path.basename(file_path),
                        "document_name": file_name,
                        "title": title,
                        "chunk_number": chunk_num + 1,
                        "text": chunk_text[:8000],
                        "char_count": len(chunk_text),
                        "file_type": "docx"
                    }

                    vectors_batch.append({
                        "id": vector_id,
                        "values": embedding,
                        "metadata": metadata
                    })
                    
                    if len(vectors_batch) >= batch_size or chunk_num == len(text_chunks) - 1:
                        if vectors_batch:
                            self.index.upsert(vectors=vectors_batch)
                            print(f"Uploaded batch of {len(vectors_batch)} vectors")
                            vectors_batch = []
            
            else:
                print(f"Unsupported file type: {file_extension}")
                return
                
            print(f"Successfully processed document: {file_name}")
            
        except Exception as e:
            print(f"Error processing document {file_path}: {e}")
            raise

    def query_similar(self, query_text: str, top_k: int = 5):
        """Query Pinecone for similar documents based on the input text."""
        query_embedding = self.create_embedding(query_text)
        if not query_embedding:
            print("Failed to create query embedding")
            return {"matches": []}
        results = self.index.query(
            vector=query_embedding,
            top_k=top_k,
            include_metadata=True,
        )
        return results



    def process_specific_document(self, filename: str, title: str, replace_existing: bool = True):
        """
        Process a specific PDF or DOCX file from the folder.
        If replace_existing is True, it will replace any existing vectors for the document with the same name.
        """
        file_extension = os.path.splitext(filename)[1].lower()
        if file_extension in [".pdf", ".docx"]:
            file_path = os.path.join(self.folder_path, filename)
            if os.path.exists(file_path):
                print(f"Processing specific file: {file_path}")
                self.embed_and_store_document(file_path, title=title, replace_existing=replace_existing)
            else:
                print(f"File not found: {file_path}")
        else:
            print(f"Unsupported file type: {filename}")


# Example usage
if __name__ == "__main__":
    vectorizer = WordVectorizerOpenAIPinecone(
        folder_path="D:\\brindyjean\\House Notes PDF",
        # folder_path="D:\\brindyjean\\Guest Messaging Guide",  # Specify the folder containing your PDFs
        pinecone_index_name="guest-messaging-guide-embedding"
    )
   
    # Example : Process a specific PDF file (replacing if it exists)
    vectorizer.process_specific_document("Guest Messaging Guide (Brandyjeans).pdf",
                                     title = "Guest Messaging Guide",
                                     replace_existing=True)